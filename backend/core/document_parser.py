import io
import pymupdf as fitz
import docx
import pandas as pd
from typing import List, Dict, Any, Tuple

class DocumentParser:
    """Extracts raw text and metadata from PDF, DOCX, CSV, XLSX, and TXT files."""

    @staticmethod
    def parse_file(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Parses uploaded file bytes based on filename extension.
        Returns a list of extracted document sections with text and metadata.
        """
        ext = filename.lower().split(".")[-1]
        
        if ext == "pdf":
            return DocumentParser.parse_pdf(file_bytes, filename)
        elif ext in ["docx", "doc"]:
            return DocumentParser.parse_docx(file_bytes, filename)
        elif ext in ["xlsx", "xls"]:
            return DocumentParser.parse_excel(file_bytes, filename)
        elif ext == "csv":
            return DocumentParser.parse_csv(file_bytes, filename)
        elif ext in ["txt", "md"]:
            return DocumentParser.parse_txt(file_bytes, filename)
        else:
            raise ValueError(f"Unsupported file format extension: .{ext}")

    @staticmethod
    def parse_pdf(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
        sections = []
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        total_pages = len(doc)

        for page_num in range(total_pages):
            page = doc[page_num]
            text = page.get_text("text").strip()
            if text:
                sections.append({
                    "text": text,
                    "metadata": {
                        "filename": filename,
                        "file_type": "pdf",
                        "page_number": page_num + 1,
                        "total_pages": total_pages,
                        "section": f"Page {page_num + 1}"
                    }
                })
        doc.close()
        return sections

    @staticmethod
    def parse_docx(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
        sections = []
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
        except Exception as e:
            raise ValueError(f"Could not parse Word document '{filename}'. Note: Legacy binary .doc files are not supported, please convert to .docx: {e}")

        current_section = "Header / General"
        current_paragraphs = []
        section_idx = 1

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            
            # Check if paragraph is a heading
            style_name = getattr(p.style, "name", "") if p.style else ""
            if style_name and style_name.lower().startswith("heading"):
                if current_paragraphs:
                    combined_text = "\n".join(current_paragraphs)
                    sections.append({
                        "text": combined_text,
                        "metadata": {
                            "filename": filename,
                            "file_type": "docx",
                            "section": current_section,
                            "section_index": section_idx
                        }
                    })
                    section_idx += 1
                    current_paragraphs = []
                current_section = text
            else:
                current_paragraphs.append(text)

        if current_paragraphs:
            sections.append({
                "text": "\n".join(current_paragraphs),
                "metadata": {
                    "filename": filename,
                    "file_type": "docx",
                    "section": current_section,
                    "section_index": section_idx
                }
            })

        # Also extract tables from docx
        for t_idx, table in enumerate(doc.tables, 1):
            table_data = []
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    table_data.append(row_text)
            if table_data:
                sections.append({
                    "text": "\n".join(table_data),
                    "metadata": {
                        "filename": filename,
                        "file_type": "docx",
                        "section": f"Table {t_idx}",
                        "section_index": section_idx + t_idx
                    }
                })

        return sections

    @staticmethod
    def parse_excel(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
        sections = []
        excel_file = pd.ExcelFile(io.BytesIO(file_bytes))

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            df = df.dropna(how="all")
            if df.empty:
                continue

            # Detect true header row if columns contain "Unnamed"
            if any("unnamed" in str(c).lower() for c in df.columns):
                for i in range(min(6, len(df))):
                    row_vals = [str(x).strip() for x in df.iloc[i].values if pd.notna(x)]
                    if any("sem" in x.lower() or "programme" in x.lower() or "fee" in x.lower() for x in row_vals):
                        headers = [str(x).replace("\n", " ").strip() if pd.notna(x) else f"Col_{j}" for j, x in enumerate(df.iloc[i].values)]
                        df = df.iloc[i+1:].copy()
                        df.columns = headers
                        break

            # Clean column headers
            df.columns = [str(c).replace("\n", " ").strip() for c in df.columns]
            df = df.dropna(how="all")

            cols = [str(c).strip() for c in df.columns]
            for r_idx, row in df.iterrows():
                items = []
                for col, val in zip(cols, row):
                    if pd.notna(val) and str(val).strip() and str(val).strip().lower() != "nan":
                        items.append(f"{col}: {str(val).strip()}")
                if items:
                    row_text = f"[Fee Structure Record {r_idx + 1} - Sheet: {sheet_name}] " + " | ".join(items)
                    sections.append({
                        "text": row_text,
                        "metadata": {
                            "filename": filename,
                            "file_type": "xlsx",
                            "sheet_name": sheet_name,
                            "row_index": r_idx + 1,
                            "section": f"Sheet: {sheet_name} (Row {r_idx + 1})"
                        }
                    })
        return sections

    @staticmethod
    def parse_csv(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
        sections = []
        try:
            df = pd.read_csv(io.BytesIO(file_bytes))
        except UnicodeDecodeError:
            df = pd.read_csv(io.BytesIO(file_bytes), encoding="latin-1")

        df = df.dropna(how="all")

        if not df.empty:
            markdown_table = df.to_markdown(index=False)
            
            row_sentences = []
            cols = [str(c).strip() for c in df.columns]
            for r_idx, row in df.iterrows():
                items = []
                for col, val in zip(cols, row):
                    if pd.notna(val) and str(val).strip() and str(val).strip().lower() != "nan":
                        items.append(f"{col}: {str(val).strip()}")
                if items:
                    row_sentences.append(f"[Record {r_idx + 1}] " + " | ".join(items))

            combined_text = f"CSV Data Table:\n{markdown_table}\n\n### Detailed Records:\n" + "\n".join(row_sentences)

            sections.append({
                "text": combined_text,
                "metadata": {
                    "filename": filename,
                    "file_type": "csv",
                    "total_rows": len(df),
                    "section": "CSV Data Table"
                }
            })
        return sections

    @staticmethod
    def parse_txt(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
        try:
            content = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            content = file_bytes.decode("latin-1", errors="ignore")

        content = content.strip()
        if not content:
            return []

        return [{
            "text": content,
            "metadata": {
                "filename": filename,
                "file_type": filename.lower().split(".")[-1],
                "section": "Document Content"
            }
        }]
